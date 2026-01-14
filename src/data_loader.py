from pathlib import Path
from typing import List, Any
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader
from langchain_community.document_loaders.excel import UnstructuredExcelLoader
from langchain_community.document_loaders import JSONLoader
from langchain_community.document_loaders import UnstructuredPowerPointLoader
from langchain_community.document_loaders import UnstructuredHTMLLoader
from langchain_community.document_loaders import UnstructuredMarkdownLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
import pytesseract
from PIL import Image

def load_all_documents(data_dir: str) -> List[Any]:
    """
    Load all supported files from the data directory and convert to LangChain document structure.
    Supported: PDF, TXT, CSV, Excel, Word, JSON, PowerPoint, HTML, Markdown, Images (with OCR)
    """
    # Use project root data folder
    data_path = Path(data_dir).resolve()
    print(f"[DEBUG] Data path: {data_path}")
    documents = []

    # PDF files
    pdf_files = list(data_path.glob('**/*.pdf'))
    print(f"[DEBUG] Found {len(pdf_files)} PDF files: {[str(f) for f in pdf_files]}")
    for pdf_file in pdf_files:
        print(f"[DEBUG] Loading PDF: {pdf_file}")
        try:
            loader = PyPDFLoader(str(pdf_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} PDF docs from {pdf_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load PDF {pdf_file}: {e}")

    # TXT files
    txt_files = list(data_path.glob('**/*.txt'))
    print(f"[DEBUG] Found {len(txt_files)} TXT files: {[str(f) for f in txt_files]}")
    for txt_file in txt_files:
        print(f"[DEBUG] Loading TXT: {txt_file}")
        try:
            loader = TextLoader(str(txt_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} TXT docs from {txt_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load TXT {txt_file}: {e}")

    # CSV files
    csv_files = list(data_path.glob('**/*.csv'))
    print(f"[DEBUG] Found {len(csv_files)} CSV files: {[str(f) for f in csv_files]}")
    for csv_file in csv_files:
        print(f"[DEBUG] Loading CSV: {csv_file}")
        try:
            loader = CSVLoader(str(csv_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} CSV docs from {csv_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load CSV {csv_file}: {e}")

    # Excel files
    xlsx_files = list(data_path.glob('**/*.xlsx'))
    print(f"[DEBUG] Found {len(xlsx_files)} Excel files: {[str(f) for f in xlsx_files]}")
    for xlsx_file in xlsx_files:
        print(f"[DEBUG] Loading Excel: {xlsx_file}")
        try:
            loader = UnstructuredExcelLoader(str(xlsx_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} Excel docs from {xlsx_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load Excel {xlsx_file}: {e}")

    # Word files
    docx_files = list(data_path.glob('**/*.docx'))
    print(f"[DEBUG] Found {len(docx_files)} Word files: {[str(f) for f in docx_files]}")
    for docx_file in docx_files:
        print(f"[DEBUG] Loading Word: {docx_file}")
        try:
            doc = DocxDocument(str(docx_file))
            text = "\n".join([para.text for para in doc.paragraphs])
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": str(docx_file), "type": "docx"}
                ))
                print(f"[DEBUG] Loaded Word doc from {docx_file}")
            else:
                print(f"[WARNING] Word file {docx_file} has no text content")
        except Exception as e:
            print(f"[ERROR] Failed to load Word {docx_file}: {e}")

    # JSON files
    json_files = list(data_path.glob('**/*.json'))
    print(f"[DEBUG] Found {len(json_files)} JSON files: {[str(f) for f in json_files]}")
    for json_file in json_files:
        print(f"[DEBUG] Loading JSON: {json_file}")
        try:
            loader = JSONLoader(str(json_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} JSON docs from {json_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load JSON {json_file}: {e}")

    # PowerPoint files
    pptx_files = list(data_path.glob('**/*.pptx')) + list(data_path.glob('**/*.ppt'))
    print(f"[DEBUG] Found {len(pptx_files)} PowerPoint files: {[str(f) for f in pptx_files]}")
    for pptx_file in pptx_files:
        print(f"[DEBUG] Loading PowerPoint: {pptx_file}")
        try:
            loader = UnstructuredPowerPointLoader(str(pptx_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} PowerPoint docs from {pptx_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load PowerPoint {pptx_file}: {e}")

    # HTML files
    html_files = list(data_path.glob('**/*.html')) + list(data_path.glob('**/*.htm'))
    print(f"[DEBUG] Found {len(html_files)} HTML files: {[str(f) for f in html_files]}")
    for html_file in html_files:
        print(f"[DEBUG] Loading HTML: {html_file}")
        try:
            loader = UnstructuredHTMLLoader(str(html_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} HTML docs from {html_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load HTML {html_file}: {e}")

    # Markdown files
    md_files = list(data_path.glob('**/*.md'))
    print(f"[DEBUG] Found {len(md_files)} Markdown files: {[str(f) for f in md_files]}")
    for md_file in md_files:
        print(f"[DEBUG] Loading Markdown: {md_file}")
        try:
            loader = UnstructuredMarkdownLoader(str(md_file))
            loaded = loader.load()
            print(f"[DEBUG] Loaded {len(loaded)} Markdown docs from {md_file}")
            documents.extend(loaded)
        except Exception as e:
            print(f"[ERROR] Failed to load Markdown {md_file}: {e}")

    # Image files with OCR
    image_extensions = ['**/*.png', '**/*.jpg', '**/*.jpeg', '**/*.bmp', '**/*.gif', '**/*.tiff']
    image_files = []
    for ext in image_extensions:
        image_files.extend(list(data_path.glob(ext)))
    print(f"[DEBUG] Found {len(image_files)} Image files: {[str(f) for f in image_files]}")
    
    for image_file in image_files:
        print(f"[DEBUG] Loading Image with OCR: {image_file}")
        try:
            img = Image.open(str(image_file))
            text = pytesseract.image_to_string(img)
            if text.strip():
                doc = Document(page_content=text, metadata={"source": str(image_file)})
                documents.append(doc)
                print(f"[DEBUG] Extracted text from image: {image_file}")
            else:
                print(f"[DEBUG] No text found in image: {image_file}")
        except Exception as e:
            print(f"[ERROR] Failed to load Image {image_file}: {e}")

    print(f"[DEBUG] Total loaded documents: {len(documents)}")
    return documents

# Example usage
if __name__ == "__main__":
    docs = load_all_documents("data")
    print(f"Loaded {len(docs)} documents.")
    print("Example document:", docs[0] if docs else None)